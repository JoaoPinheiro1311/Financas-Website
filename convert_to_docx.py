import sys
from docx import Document
from docx.shared import Pt
import os

def markdown_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    doc = Document()
    
    # Adicionar título principal se existir no início do MD
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif '|' in line and '--' not in line:
            # Simples tratamento de tabelas (apenas texto por enquanto)
            doc.add_paragraph(line.replace('|', '  '), style='Normal')
        elif line.startswith('```'):
            # Ignorar blocos de código/mermaid por agora ou tratar como texto simples
            continue
        else:
            # Texto normal
            p = doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"Success: {docx_path} created")

if __name__ == "__main__":
    # Relatório
    md_file = r"C:\Users\João Pinheiro\.gemini\antigravity\brain\c2a9de1e-bb2f-4dac-abec-10dbc5e3ac74\RELATORIO.md"
    docx_file = r"c:\Users\João Pinheiro\Desktop\Financas-Website\RELATORIO.docx"
    markdown_to_docx(md_file, docx_file)
    
    # Walkthrough
    md_walk = r"C:\Users\João Pinheiro\.gemini\antigravity\brain\c2a9de1e-bb2f-4dac-abec-10dbc5e3ac74\walkthrough.md"
    docx_walk = r"c:\Users\João Pinheiro\Desktop\Financas-Website\walkthrough.docx"
    markdown_to_docx(md_walk, docx_walk)
